import sys
import io

# Forzar codificación UTF-8 en stdout
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
from docxcompose.composer import Composer 
from docx import Document

class ReporteDocumento:
    def __init__(self, categoria):
        self.categoria = categoria
        self.documento = Document()
        self.nombre_archivo = f"reporte_{categoria}.docx"  #sin timestamp

        #Determina ruta y logo
        if self.nombre_archivo.startswith(("reporte_Login", "reporte_actualizar_datos_WFM")):
            self.ruta_carpeta = os.path.abspath(os.path.join("../reports", "word","all_reports"))
            logo_path = "../utils/logotipos/Claro.png"
        else:
            self.ruta_carpeta = os.path.abspath(os.path.join("reports", "word","all_reports"))
            logo_path = "utils/logotipos/Claro.png"

        #Crea carpeta si no existe
        if not os.path.exists(self.ruta_carpeta):
            os.makedirs(self.ruta_carpeta)

        #Borra archivo anterior (solo uno por categoría)
        ruta_archivo_anterior = os.path.join(self.ruta_carpeta, self.nombre_archivo)
        if os.path.exists(ruta_archivo_anterior):
            try:
                os.remove(ruta_archivo_anterior)
            except Exception as e:
                print(f"No se pudo eliminar el reporte anterior: {e}")

        # Header con logo y categoría
        for section in self.documento.sections:
            header = section.header
            table = header.add_table(rows=1, cols=2, width=Inches(6.5))
            table.allow_autofit = True
            table.columns[0].width = Inches(5)
            table.columns[1].width = Inches(1.5)

            cell_text = table.cell(0, 0)
            p_text = cell_text.paragraphs[0]
            run = p_text.add_run(f"\nPruebas Módulo de Gestión\nCategoría {self.categoria}")
            run.font.size = Pt(13)
            p_text.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if os.path.exists(logo_path):
                cell_img = table.cell(0, 1)
                p_img = cell_img.paragraphs[0]
                run = p_img.add_run()
                run.add_picture(logo_path, width=Inches(1.0))
                p_img.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                print(f"No se encontró el logo en: {logo_path}")

        self.es_primera_evidencia = True

    def agregar_evidencia(self, nombre_prueba, ruta_imagen):
        if not self.es_primera_evidencia:
            self.documento.add_page_break()

        self.documento.add_paragraph().add_run("Fecha de ejecución: ").bold = True
        self.documento.paragraphs[-1].add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

        self.documento.add_paragraph().add_run("Categoría de prueba: ").bold = True
        self.documento.paragraphs[-1].add_run(self.categoria)

        self.documento.add_paragraph().add_run("Nombre de prueba: ").bold = True
        self.documento.paragraphs[-1].add_run(nombre_prueba)

        self.documento.add_paragraph().add_run("Evidencia:").bold = True

        if os.path.exists(ruta_imagen):
            p_img = self.documento.add_paragraph()
            p_img.add_run().add_picture(ruta_imagen, width=Inches(5.5))
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            self.documento.add_paragraph("[Imagen no encontrada]")

        self.es_primera_evidencia = False

    
    def guardar(self):
        ruta_final = os.path.join(self.ruta_carpeta, self.nombre_archivo)

        try:
            os.makedirs(os.path.dirname(ruta_final), exist_ok=True)
            self.documento.save(ruta_final)
            print(f"Reporte Word guardado en: {ruta_final}")
            return os.path.abspath(ruta_final)
        except PermissionError:
            print("El archivo está abierto. Por favor, ciérrelo e intente de nuevo.")
            return None
        except Exception as e:
            print(f"Error inesperado al guardar el Word: {e}")
            return None

# Ruta donde están los documentos .docx
carpeta_actual = os.path.dirname(os.path.abspath(__file__))
carpeta_reportes = os.path.join(carpeta_actual, '..', '..', 'reports', 'word', 'all_reports')
carpeta_reportes = os.path.abspath(carpeta_reportes)

# Verificar que la carpeta existe
if not os.path.exists(carpeta_reportes):
    print(f"[WARNING] Carpeta de reportes no encontrada: {carpeta_reportes}")
else:
    documentos = [f for f in os.listdir(carpeta_reportes) if f.endswith(".docx")]
    
    if not documentos:
        print("[WARNING] No se encontraron documentos .docx para unificar")
    else:
        # Crear documento maestro vacío
        from docx import Document
        documento_unificado = Document()

        # === CABECERA PARA DOCUMENTO UNIFICADO ===
        # Determinar ruta del logo (misma lógica que en ReporteDocumento)
        logo_path = os.path.abspath(os.path.join(carpeta_actual, '..', 'logotipos', 'Claro.png'))
        if not os.path.exists(logo_path):
            logo_path = os.path.abspath(os.path.join(carpeta_actual, '..', '..', 'utils', 'logotipos', 'Claro.png'))
        
        # Obtener fecha y hora actual
        fecha_actual = datetime.now().strftime('%Y-%m-%d')
        hora_actual = datetime.now().strftime('%H:%M:%S')

        # Crear cabecera
        for section in documento_unificado.sections:
            header = section.header
            table = header.add_table(rows=1, cols=2, width=Inches(6.5))
            table.allow_autofit = True
            table.columns[0].width = Inches(5)
            table.columns[1].width = Inches(1.5)

            cell_text = table.cell(0, 0)
            p_text = cell_text.paragraphs[0]
            run = p_text.add_run(f"\nREPORTE GENERAL DE PRUEBAS\nFecha de generación: {fecha_actual}\Hora de generación: {hora_actual}")
            run.font.size = Pt(12)
            p_text.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if os.path.exists(logo_path):
                cell_img = table.cell(0, 1)
                p_img = cell_img.paragraphs[0]
                run = p_img.add_run()
                run.add_picture(logo_path, width=Inches(1.0))
                p_img.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                print(f"[WARNING] No se encontró el logo en: {logo_path}")

        # Añadir título principal
        #titulo = documento_unificado.add_paragraph()
        #titulo.add_run("REPORTE GENERAL DE PRUEBAS").bold = True
        #titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #titulo.runs[0].font.size = Pt(16)
        
        # Añadir fecha
        #fecha_parrafo = documento_unificado.add_paragraph()
        #from datetime import datetime
        #fecha_parrafo.add_run(f"Fecha de generación: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        #fecha_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Añadir un párrafo vacío para evitar problemas con docxcompose
        documento_unificado.add_paragraph("")

        # Componer todos los documentos
        composer = Composer(documento_unificado)

        # Ordenar documentos para consistencia
        documentos.sort()
        
        for docx_file in documentos:
            try:
                ruta_docx = os.path.join(carpeta_reportes, docx_file)
                #print(f"[INFO] Procesando: {docx_file}")
                
                # Cargar documento
                doc = Document(ruta_docx)
                composer.append(doc)
                
            except Exception as e:
                print(f"[ERROR] No se pudo procesar {docx_file}: {e}")

        # Guardar el documento final
        ruta_salida = os.path.abspath(os.path.join(
            carpeta_actual, '..', '..', 'reports', 'word', 'one_report_general', 'documento_unificado.docx'
        ))
        
        # Crear directorio si no existe
        os.makedirs(os.path.dirname(ruta_salida), exist_ok=True)
        
        try:
            composer.save(ruta_salida)
            print(f"[OK] Documentos unificados correctamente en: {ruta_salida}")
        except Exception as e:
            print(f"[ERROR] No se pudo guardar el documento unificado: {e}")